from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import IO

import httpx
from bs4 import BeautifulSoup

# pdfplumber and python-docx are heavyweight optional deps. Import lazily so a
# system without them can still boot — the endpoint reports a clean 415 for the
# affected content type instead of crashing at import time.

_MAX_FETCH_BYTES = 10 * 1024 * 1024  # 10 MB ceiling on URL/file payloads
_USER_AGENT = "fed-pulse-document-parser/1.0"
_WHITESPACE_RE = re.compile(r"[ \t ]+")
_NEWLINE_RE = re.compile(r"\n{3,}")


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    char_count: int
    source_kind: str
    source_metadata: dict[str, str]

    def to_dict(self) -> dict[str, object]:
        return {
            "text": self.text,
            "char_count": self.char_count,
            "source_kind": self.source_kind,
            "source_metadata": dict(self.source_metadata),
        }


def normalise(text: str) -> str:
    """Collapse runs of internal whitespace and triple-blank-lines into one
    canonical shape so all three ingestion modes produce the same input to
    the forecaster."""

    if not text:
        return ""
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = _WHITESPACE_RE.sub(" ", cleaned)
    cleaned = _NEWLINE_RE.sub("\n\n", cleaned)
    return cleaned.strip()


def parse_paste(raw: str) -> ParsedDocument:
    text = normalise(raw)
    return ParsedDocument(
        text=text,
        char_count=len(text),
        source_kind="paste",
        source_metadata={},
    )


def parse_pdf_stream(stream: IO[bytes], *, filename: str | None = None) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pdfplumber not installed; cannot parse PDF uploads") from exc

    pages: list[str] = []
    with pdfplumber.open(stream) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text:
                pages.append(text)
    body = normalise("\n\n".join(pages))
    return ParsedDocument(
        text=body,
        char_count=len(body),
        source_kind="pdf",
        source_metadata={"filename": filename or "", "page_count": str(len(pages))},
    )


def parse_docx_stream(stream: IO[bytes], *, filename: str | None = None) -> ParsedDocument:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx not installed; cannot parse DOCX uploads") from exc

    doc = docx.Document(stream)
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    body = normalise("\n\n".join(paragraphs))
    return ParsedDocument(
        text=body,
        char_count=len(body),
        source_kind="docx",
        source_metadata={"filename": filename or "", "paragraph_count": str(len(paragraphs))},
    )


def _extract_visible_text(html: str) -> str:
    """Pull article-class or main-tag visible text out of an HTML payload.

    The strategy is intentionally conservative: prefer `<article>` or `<main>`
    if present, fall back to body text with script/style stripped.
    """

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "nav", "footer", "header", "aside"]):
        tag.decompose()
    candidate = soup.find("article") or soup.find("main") or soup.body or soup
    text = candidate.get_text(separator="\n") if candidate else ""
    return normalise(text)


async def parse_url(url: str, *, client: httpx.AsyncClient | None = None) -> ParsedDocument:
    headers = {"User-Agent": _USER_AGENT, "Accept": "text/html,application/pdf"}
    owns_client = client is None
    if owns_client:
        client = httpx.AsyncClient(headers=headers, follow_redirects=True, timeout=20.0)
    try:
        response = await client.get(url)
        response.raise_for_status()
        if int(response.headers.get("content-length", "0")) > _MAX_FETCH_BYTES:
            raise ValueError(f"Response too large for inline parsing (>{_MAX_FETCH_BYTES} bytes)")
        content_type = response.headers.get("content-type", "").split(";")[0].strip().lower()
        if content_type == "application/pdf":
            parsed = parse_pdf_stream(io.BytesIO(response.content), filename=url.rsplit("/", 1)[-1])
            return ParsedDocument(
                text=parsed.text,
                char_count=parsed.char_count,
                source_kind="url",
                source_metadata={
                    **parsed.source_metadata,
                    "url": url,
                    "content_type": content_type,
                },
            )
        text = _extract_visible_text(response.text)
        return ParsedDocument(
            text=text,
            char_count=len(text),
            source_kind="url",
            source_metadata={
                "url": url,
                "content_type": content_type or "text/html",
            },
        )
    finally:
        if owns_client and client is not None:
            await client.aclose()
